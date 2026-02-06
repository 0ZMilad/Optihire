"""
Resume parsing and background processing service.

This module handles:
- Downloading resume files from Supabase storage
- Parsing PDF and DOCX files to extract text
- Extracting structured data (contact info, experience, education, skills)
- Background processing with status updates
"""
import io
import re
from uuid import UUID
from datetime import datetime, date
from typing import Any, Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pdfminer.high_level import extract_text as extract_pdf_text
from pdfminer.pdfparser import PDFSyntaxError
from pdfminer.pdfdocument import PDFEncryptionError
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from sqlmodel import Session, select
from sqlalchemy.exc import IntegrityError

from app.core.config import settings
from app.core.logging_config import log_info, log_error, log_warning
from app.db.session import SessionLocal
from app.models.resume_model import (
    Resume,
    ResumeSkill,
    ResumeExperience,
    ResumeEducation,
    ResumeCertification,
    ResumeProject,
)
from app.schemas.resume_schema import ResumeParseStatusResponse, ResumeCreate
from app.services.storage_service import get_supabase_client

# =============================================================================
# RESUME CREATION
# =============================================================================

def create_resume(db: Session, resume_data: ResumeCreate) -> Resume:
    """
    Create a new resume with manual data input.
    
    Args:
        db: Database session
        resume_data: Resume creation data
        
    Returns:
        Created Resume instance
        
    Raises:
        IntegrityError: If version name already exists for the user
    """
    # Check if version name already exists for this user
    existing_resume = db.exec(
        select(Resume).where(
            Resume.user_id == resume_data.user_id,
            Resume.version_name == resume_data.version_name,
            Resume.deleted_at.is_(None)
        )
    ).first()
    
    if existing_resume:
        raise IntegrityError(
            "Version name already exists for this user",
            None,
            None
        )
    
    # Create the resume with "Completed" status since it's manually created
    resume = Resume(
        user_id=resume_data.user_id,
        version_name=resume_data.version_name,
        template_id=resume_data.template_id,
        is_primary=resume_data.is_primary,
        full_name=resume_data.full_name,
        email=resume_data.email,
        phone=resume_data.phone,
        location=resume_data.location,
        linkedin_url=resume_data.linkedin_url,
        github_url=resume_data.github_url,
        portfolio_url=resume_data.portfolio_url,
        professional_summary=resume_data.professional_summary,
        processing_status="Completed",  # Manual creation is immediately complete
        last_analyzed_at=datetime.utcnow()
    )
    
    db.add(resume)
    db.commit()
    db.refresh(resume)
    
    log_info(f"Created new resume manually: user_id={resume_data.user_id}, version_name={resume_data.version_name}")
    
    return resume


# =============================================================================
# STATUS MANAGEMENT
# =============================================================================

def get_parse_status(
    resume_id: UUID, 
    user_id: UUID, 
    db: Session
) -> Optional[ResumeParseStatusResponse]:
    """
    Retrieve the current parsing status of a resume.
    
    Args:
        resume_id: UUID of the resume to check
        user_id: UUID of the user (for authorization)
        db: Database session
        
    Returns:
        ResumeParseStatusResponse if resume found and user authorized, None otherwise
    """
    # Query the resume with ownership check
    statement = select(Resume).where(
        Resume.id == resume_id,
        Resume.user_id == user_id
    )
    result = db.exec(statement)
    resume = result.first()
    
    if not resume:
        return None
    
    # Map status to user-friendly messages
    status_messages = {
        "Pending": "Resume is queued for parsing",
        "Processing": "Parsing resume content...",
        "Completed": "Resume parsed successfully",
        "Failed": "Failed to parse resume"
    }
    
    message = status_messages.get(resume.processing_status, "Unknown status")
    
    return ResumeParseStatusResponse(
        id=resume.id,
        status=resume.processing_status,
        message=message,
        created_at=resume.created_at,
        updated_at=resume.updated_at,
        error_details=resume.error_message if resume.processing_status == "Failed" else None
    )


# =============================================================================
# CONSTANTS
# =============================================================================

# Minimum character threshold for scanned PDF detection (Task 4555)
MIN_TEXT_LENGTH_THRESHOLD = 50

# Max file size
MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024

# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

# Regex patterns for structured data extraction
PATTERNS = {
    # Email pattern
    "email": re.compile(
        r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
        re.IGNORECASE
    ),
    # Phone patterns (international and US formats)
    "phone": re.compile(
        r"(?:\+?1[-.\s]?)?"  # Optional country code
        r"(?:\(?\d{3}\)?[-.\s]?)"  # Area code
        r"\d{3}[-.\s]?\d{4}"  # Main number
        r"|\+\d{1,3}[-.\s]?\d{1,14}",  # International format
        re.IGNORECASE
    ),
    # LinkedIn URL
    "linkedin": re.compile(
        r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+/?",
        re.IGNORECASE
    ),
    # GitHub URL
    "github": re.compile(
        r"(?:https?://)?(?:www\.)?github\.com/[\w-]+/?",
        re.IGNORECASE
    ),
    # Portfolio/website URL
    "website": re.compile(
        r"(?:https?://)?(?:www\.)?[\w-]+\.(?:com|io|dev|me|org|net)(?:/[\w-]*)*",
        re.IGNORECASE
    ),
}

# Section header patterns for resume sectioning
# These patterns must match section headers at line boundaries, not words in prose
SECTION_HEADERS = {
    "experience": re.compile(
        r"^\s*(?:work\s+)?experience(?:\s+history)?(?:\s*:|\s*$)|^\s*employment(?:\s+history)?(?:\s*:|\s*$)|^\s*professional\s+(?:experience|background)(?:\s*:|\s*$)|^\s*career\s+history(?:\s*:|\s*$)",
        re.IGNORECASE | re.MULTILINE
    ),
    "education": re.compile(
        r"^\s*education(?:al\s+background)?(?:\s*:|\s*$)|^\s*academic(?:\s+background)?(?:\s*:|\s*$)|^\s*qualifications(?:\s*:|\s*$)|^\s*degrees?(?:\s*:|\s*$)",
        re.IGNORECASE | re.MULTILINE
    ),
    "skills": re.compile(
        r"^\s*(?:technical\s+)?skills(?:\s*:|\s*$)|^\s*competenc(?:ies|e)(?:\s*:|\s*$)|^\s*expertise(?:\s*:|\s*$)|^\s*technologies(?:\s*:|\s*$)|^\s*proficiencies(?:\s*:|\s*$)",
        re.IGNORECASE | re.MULTILINE
    ),
    "certifications": re.compile(
        r"^\s*certifications?(?:\s*:|\s*$)|^\s*licenses?(?:\s*:|\s*$)|^\s*credentials?(?:\s*:|\s*$)|^\s*professional\s+development(?:\s*:|\s*$)",
        re.IGNORECASE | re.MULTILINE
    ),
    "projects": re.compile(
        r"^\s*projects?(?:\s*:|\s*$)|^\s*portfolio(?:\s*:|\s*$)|^\s*personal\s+projects?(?:\s*:|\s*$)|^\s*side\s+projects?(?:\s*:|\s*$)",
        re.IGNORECASE | re.MULTILINE
    ),
    "summary": re.compile(
        r"^\s*(?:professional\s+)?summary(?:\s*:|\s*$)|^\s*profile(?:\s*:|\s*$)|^\s*objective(?:\s*:|\s*$)|^\s*about(?:\s+me)?(?:\s*:|\s*$)|^\s*overview(?:\s*:|\s*$)",
        re.IGNORECASE | re.MULTILINE
    ),
}


# =============================================================================
# DATABASE OPERATIONS
# =============================================================================


def update_resume_status(
    resume_id: UUID,
    status: str,
    error_message: str | None = None,
    db: Session | None = None
) -> None:
    """
    Update the processing status of a resume.
    
    Args:
        resume_id: The UUID of the resume to update
        status: The new status ('Pending', 'Processing', 'Completed', 'Failed')
        error_message: Optional error message if status is 'Failed'
        db: Optional database session (will create one if not provided)
    """
    close_db = False
    if db is None:
        db = SessionLocal()
        close_db = True
    
    try:
        statement = select(Resume).where(Resume.id == resume_id)
        resume = db.exec(statement).first()
        
        if resume:
            resume.processing_status = status
            resume.error_message = error_message
            
            if status == "Completed":
                resume.last_analyzed_at = datetime.utcnow()
            
            db.add(resume)
            db.commit()
            log_info(f"Resume {resume_id} status updated to {status}")
        else:
            log_warning(f"Resume {resume_id} not found for status update")
    
    except Exception as e:
        log_error(f"Failed to update resume status for {resume_id}: {str(e)}")
        if db:
            db.rollback()
    finally:
        if close_db and db:
            db.close()


def update_resume_with_parsed_data(
    resume: Resume,
    parsed_data: dict[str, Any],
    db: Session
) -> None:
    """
    Update resume record with extracted/parsed data.
    
    This function persists both flat contact fields and structured data
    (skills, experiences, education) to their respective normalized tables.
    
    Args:
        resume: The Resume model instance to update
        parsed_data: Dictionary containing extracted resume information
        db: Database session
    """
    # Update contact information
    if parsed_data.get("full_name"):
        resume.full_name = parsed_data["full_name"]
    if parsed_data.get("email"):
        resume.email = parsed_data["email"]
    if parsed_data.get("phone"):
        resume.phone = parsed_data["phone"]
    if parsed_data.get("location"):
        resume.location = parsed_data["location"]
    if parsed_data.get("linkedin_url"):
        resume.linkedin_url = parsed_data["linkedin_url"]
    if parsed_data.get("github_url"):
        resume.github_url = parsed_data["github_url"]
    if parsed_data.get("portfolio_url"):
        resume.portfolio_url = parsed_data["portfolio_url"]
    if parsed_data.get("professional_summary"):
        resume.professional_summary = parsed_data["professional_summary"]
    
    # Store raw text for future reprocessing
    if parsed_data.get("raw_text"):
        resume.raw_text = parsed_data["raw_text"]
        log_info(f"Stored {len(parsed_data['raw_text'])} characters of raw text for resume {resume.id}")
    
    db.add(resume)
    
    # ==========================================================================
    # PERSIST STRUCTURED DATA TO NORMALIZED TABLES
    # ==========================================================================
    
    # Skills - persist to resume_skills table
    skills_list = parsed_data.get("skills", [])
    if skills_list:
        log_info(f"Persisting {len(skills_list)} skills for resume {resume.id}")
        for i, skill_name in enumerate(skills_list):
            if isinstance(skill_name, str) and skill_name.strip():
                skill = ResumeSkill(
                    resume_id=resume.id,
                    skill_name=skill_name.strip()[:100],  # Enforce max length
                    display_order=i
                )
                db.add(skill)
    
    # Experiences - persist to resume_experiences table
    experiences_list = parsed_data.get("experiences", [])
    if experiences_list:
        log_info(f"Persisting {len(experiences_list)} experiences for resume {resume.id}")
        for i, exp in enumerate(experiences_list):
            if isinstance(exp, dict):
                experience = ResumeExperience(
                    resume_id=resume.id,
                    company_name=exp.get("company_name") or "Unknown Company",
                    job_title=exp.get("job_title") or "Position",
                    description=exp.get("raw_text") or exp.get("description"),
                    start_date=exp.get("start_date") or date.today(),
                    end_date=exp.get("end_date"),
                    is_current=exp.get("is_current", False),
                    display_order=i
                )
                db.add(experience)
    
    # Education - persist to resume_education table
    education_list = parsed_data.get("education", [])
    if education_list:
        log_info(f"Persisting {len(education_list)} education entries for resume {resume.id}")
        for i, edu in enumerate(education_list):
            if isinstance(edu, dict):
                education = ResumeEducation(
                    resume_id=resume.id,
                    institution_name=edu.get("institution_name") or "Unknown Institution",
                    degree_type=edu.get("degree_type"),
                    field_of_study=edu.get("field_of_study"),
                    start_date=edu.get("start_date"),
                    end_date=edu.get("end_date"),
                    is_current=edu.get("is_current", False),
                    display_order=i
                )
                db.add(education)
    
    # Certifications - persist to resume_certifications table
    certifications_list = parsed_data.get("certifications", [])
    if certifications_list:
        log_info(f"Persisting {len(certifications_list)} certifications for resume {resume.id}")
        for i, cert in enumerate(certifications_list):
            if isinstance(cert, dict):
                certification = ResumeCertification(
                    resume_id=resume.id,
                    certification_name=cert.get("certification_name") or cert.get("raw_text", "Certification")[:200],
                    issuing_organization=cert.get("issuing_organization"),
                    issue_date=cert.get("issue_date"),
                    expiry_date=cert.get("expiry_date"),
                    display_order=i
                )
                db.add(certification)
    
    # Projects - persist to resume_projects table
    projects_list = parsed_data.get("projects", [])
    if projects_list:
        log_info(f"Persisting {len(projects_list)} projects for resume {resume.id}")
        for i, proj in enumerate(projects_list):
            if isinstance(proj, dict):
                project = ResumeProject(
                    resume_id=resume.id,
                    project_name=proj.get("project_name") or "Project",
                    description=proj.get("raw_text") or proj.get("description"),
                    display_order=i
                )
                db.add(project)
    
    # Single atomic commit for all changes
    db.commit()
    
    log_info(
        f"Resume {resume.id} updated with parsed data: "
        f"skills={len(skills_list)}, experiences={len(experiences_list)}, "
        f"education={len(education_list)}"
    )


# =============================================================================
# FILE DOWNLOAD
# =============================================================================

def download_resume_file(file_path: str) -> bytes:
    """
    Download resume file from Supabase storage into memory.
    
    Args:
        file_path: The storage path of the file (e.g., "resumes/uuid.pdf")
        
    Returns:
        bytes: The file content as bytes
        
    Raises:
        ValueError: With specific error codes:
            - ParseError: FileNotFound - File not in storage
            - ParseError: Oversize - File exceeds 5MB limit
            - ParseError: DownloadFailed - Storage download error
    """
    client = get_supabase_client()
    bucket = settings.SUPABASE_STORAGE_BUCKET
    
    try:
        log_info(f"Downloading file from storage: {file_path}")
        
        # Download file content as bytes
        response = client.storage.from_(bucket).download(file_path)
        
        if response is None:
            raise ValueError(f"ParseError: FileNotFound - {file_path}")
        
        # Check file size (5MB limit)
        file_size = len(response)
        if file_size > MAX_FILE_SIZE_BYTES:
            log_error(
                f"File {file_path} exceeds size limit: {file_size} bytes "
                f"(max: {MAX_FILE_SIZE_BYTES} bytes)"
            )
            raise ValueError("ParseError: Oversize")
        
        log_info(f"Successfully downloaded file: {file_path} ({file_size} bytes)")
        return response
        
    except Exception as e:
        error_msg = str(e)
        log_error(f"Failed to download file {file_path}: {error_msg}")
        
        # Check for specific error types
        if "not found" in error_msg.lower() or "404" in error_msg:
            raise ValueError(f"ParseError: FileNotFound - {file_path}")
        
        raise ValueError(f"ParseError: DownloadFailed - {error_msg}")


# =============================================================================
# PDF PARSING
# =============================================================================

def parse_pdf_resume(file_content: bytes) -> str:
    """
    Parse PDF resume and extract text content using pdfminer.six.
    
    Args:
        file_content: PDF file content as bytes
        
    Returns:
        str: Extracted text from the PDF
        
    Raises:
        ValueError: With specific error codes for different failure scenarios:
            - ParseError: ScannedPdfNoText - PDF is scanned/image-based
            - ParseError: EncryptedPdf - PDF is password protected
            - ParseError: CorruptedPdf - PDF file is corrupted/malformed
    """
    try:
        log_info("Parsing PDF document...")
        
        # Create a file-like object from bytes
        pdf_file = io.BytesIO(file_content)
        
        # Extract text using pdfminer
        extracted_text = extract_pdf_text(pdf_file)
        
        # Clean up the extracted text
        stripped_text = extracted_text.strip() if extracted_text else ""
        
        # Task 4555: Scanned PDF Detection
        # If text is too short, it's likely a scanned/image-based PDF
        if len(stripped_text) < MIN_TEXT_LENGTH_THRESHOLD:
            log_warning(
                f"PDF appears to be scanned - extracted only {len(stripped_text)} characters"
            )
            raise ValueError("ParseError: ScannedPdfNoText")
        
        log_info(f"Successfully extracted {len(stripped_text)} characters from PDF")
        return stripped_text
        
    except PDFEncryptionError:
        log_error("PDF is encrypted/password protected")
        raise ValueError("ParseError: EncryptedPdf")
        
    except PDFTextExtractionNotAllowed:
        log_error("PDF text extraction is not allowed (restricted)")
        raise ValueError("ParseError: EncryptedPdf")
        
    except PDFSyntaxError as e:
        log_error(f"PDF syntax error (corrupted file): {str(e)}")
        raise ValueError("ParseError: CorruptedPdf")
        
    except ValueError:
        # Re-raise ValueError (our custom errors) without wrapping
        raise
        
    except Exception as e:
        log_error(f"Unexpected error parsing PDF: {str(e)}")
        # Check if it's related to encryption or corruption
        error_msg = str(e).lower()
        if "encrypt" in error_msg or "password" in error_msg:
            raise ValueError("ParseError: EncryptedPdf")
        if "corrupt" in error_msg or "invalid" in error_msg or "malformed" in error_msg:
            raise ValueError("ParseError: CorruptedPdf")
        raise ValueError(f"ParseError: CorruptedPdf - {str(e)}")


# =============================================================================
# DOCX PARSING
# =============================================================================

def parse_docx_resume(file_content: bytes) -> str:
    """
    Parse DOCX resume and extract text content using python-docx.
    
    Args:
        file_content: DOCX file content as bytes
        
    Returns:
        str: Extracted text from the DOCX file
        
    Raises:
        ValueError: If the DOCX file is corrupted or invalid
    """
    try:
        log_info("Parsing DOCX document...")
        
        # Create a file-like object from bytes
        docx_file = io.BytesIO(file_content)
        
        # Load the document
        document = Document(docx_file)
        
        # Extract text from paragraphs
        paragraphs = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
        
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        extracted_text = "\n".join(paragraphs)
        
        if len(extracted_text.strip()) < MIN_TEXT_LENGTH_THRESHOLD:
            log_warning(
                f"DOCX has minimal text content - only {len(extracted_text.strip())} characters"
            )
            # For DOCX, minimal text might just be an empty/template document
            # Still allow it through but log the warning
        
        log_info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
        return extracted_text
        
    except PackageNotFoundError:
        log_error("DOCX file is corrupted or not a valid DOCX")
        raise ValueError("ParseError: CorruptedDocx")
        
    except Exception as e:
        error_msg = str(e).lower()
        log_error(f"Error parsing DOCX: {str(e)}")
        
        if "corrupt" in error_msg or "invalid" in error_msg:
            raise ValueError("ParseError: CorruptedDocx")
        
        raise ValueError(f"ParseError: CorruptedDocx - {str(e)}")


# =============================================================================
# STRUCTURED DATA EXTRACTION (Task 4268)
# =============================================================================

def _extract_contact_info(text: str) -> dict[str, str | None]:
    """
    Extract contact information from resume text.
    
    Returns:
        dict with keys: email, phone, linkedin_url, github_url, portfolio_url
    """
    contact_info: dict[str, str | None] = {
        "email": None,
        "phone": None,
        "linkedin_url": None,
        "github_url": None,
        "portfolio_url": None,
    }
    
    # Extract email
    email_match = PATTERNS["email"].search(text)
    if email_match:
        contact_info["email"] = email_match.group(0).lower()
    
    # Extract phone
    phone_match = PATTERNS["phone"].search(text)
    if phone_match:
        # Clean up phone number
        phone = re.sub(r"[^\d+]", "", phone_match.group(0))
        contact_info["phone"] = phone
    
    # Extract LinkedIn
    linkedin_match = PATTERNS["linkedin"].search(text)
    if linkedin_match:
        url = linkedin_match.group(0)
        if not url.startswith("http"):
            url = "https://" + url
        contact_info["linkedin_url"] = url
    
    # Extract GitHub
    github_match = PATTERNS["github"].search(text)
    if github_match:
        url = github_match.group(0)
        if not url.startswith("http"):
            url = "https://" + url
        contact_info["github_url"] = url
    
    return contact_info


def _extract_name(text: str) -> str | None:
    """
    Attempt to extract the candidate's name from the resume.
    
    Heuristic: The name is usually in the first few lines, 
    often the first non-empty line that looks like a name.
    """
    lines = text.strip().split("\n")
    
    for line in lines[:10]:  # Check first 10 lines
        line = line.strip()
        
        # Skip empty lines
        if not line:
            continue
        
        # Skip lines that look like headers, emails, phones, URLs
        if "@" in line or "http" in line.lower():
            continue
        if re.match(r"^[\d\s\-\(\)\+]+$", line):  # Phone numbers
            continue
        if len(line) > 50:  # Too long to be a name
            continue
        if any(keyword in line.lower() for keyword in [
            "resume", "curriculum", "cv", "objective", "summary", "experience"
        ]):
            continue
        
        # Check if it looks like a name (2-4 words, mostly letters)
        words = line.split()
        if 1 <= len(words) <= 4:
            # Check if words look like name components
            if all(re.match(r"^[A-Za-z\.\-\']+$", word) for word in words):
                return line
    
    return None


def _extract_section_content(text: str, section_name: str) -> str:
    """
    Extract content for a specific section from resume text.
    
    Args:
        text: Full resume text
        section_name: Name of the section to extract
        
    Returns:
        The content of the section (may be empty string)
    """
    if section_name not in SECTION_HEADERS:
        return ""
    
    pattern = SECTION_HEADERS[section_name]
    
    # Find where this section starts
    match = pattern.search(text)
    if not match:
        return ""
    
    section_start = match.end()
    
    # Find where the next section starts
    next_section_start = len(text)
    for other_section, other_pattern in SECTION_HEADERS.items():
        if other_section == section_name:
            continue
        other_match = other_pattern.search(text[section_start:])
        if other_match:
            potential_end = section_start + other_match.start()
            if potential_end < next_section_start:
                next_section_start = potential_end
    
    section_content = text[section_start:next_section_start].strip()
    return section_content


def _parse_skills_section(skills_text: str) -> list[str]:
    """
    Parse skills section and extract individual skills.
    """
    if not skills_text:
        return []
    
    skills = []
    
    # Split by common delimiters
    # First, try splitting by newlines
    lines = skills_text.split("\n")
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Remove category labels (e.g., "Languages:", "Frontend:", etc.)
        # Look for pattern "Word(s): content" and keep only content
        colon_match = re.match(r"^([A-Za-z\s]+):\s*(.+)$", line)
        if colon_match:
            # Take only the part after the colon
            line = colon_match.group(2)
        
        # Split by common separators: comma, semicolon, pipe, bullet points
        parts = re.split(r"[,;|•·▪◦‣⁃]\s*", line)
        
        for part in parts:
            skill = part.strip().strip("-•·").strip()
            # Filter out empty strings and very long "skills" (likely sentences)
            if skill and len(skill) < 50 and len(skill) > 1:
                skills.append(skill)
    
    # Remove duplicates while preserving order
    seen = set()
    unique_skills = []
    for skill in skills:
        skill_lower = skill.lower()
        if skill_lower not in seen:
            seen.add(skill_lower)
            unique_skills.append(skill)
    
    return unique_skills[:50]  # Limit to 50 skills


def _parse_experience_section(experience_text: str) -> list[dict[str, Any]]:
    """
    Parse experience section into structured entries.
    
    Returns a list of experience dictionaries with:
    - company_name, job_title, start_date, end_date, description, is_current
    """
    if not experience_text:
        return []
    
    experiences = []
    
    # Date pattern to identify the start of a new job entry
    # Matches formats like: 01/2015-06/2017, Jan 2015 - Present, 2015-2017, etc.
    date_pattern = re.compile(
        r"(?:"
        # MM/YYYY format: 01/2015-06/2017 or 01/2015 - Present
        r"(?:\d{1,2}/\d{4}\s*[-–—to]+\s*(?:\d{1,2}/\d{4}|[Pp]resent|[Cc]urrent))"
        r"|"
        # Month YYYY format: Jan 2015 - June 2017 or January 2015 - Present
        r"(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
        r"[\s,\.]*\d{4}\s*[-–—to]+\s*"
        r"(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
        r"[\s,\.]*\d{4}|[Pp]resent|[Cc]urrent))"
        r"|"
        # YYYY-YYYY format: 2015-2017 or 2015 - Present
        r"(?:\d{4}\s*[-–—to]+\s*(?:\d{4}|[Pp]resent|[Cc]urrent))"
        r")",
        re.IGNORECASE
    )
    
    # Company/location pattern - looks for "Company Name – City, STATE" or similar
    company_location_pattern = re.compile(
        r"^(.+?)\s*[-–—]\s*([A-Za-z\s]+,\s*[A-Z]{2})\s*$"
    )
    
    # Job title pattern - typically capitalized words at start of line
    job_title_pattern = re.compile(
        r"^([A-Z][A-Za-z\s]+(?:Manager|Engineer|Developer|Analyst|Director|"
        r"Coordinator|Specialist|Associate|Intern|Lead|Senior|Junior|"
        r"Assistant|Supervisor|Executive|Officer|Administrator|Consultant|"
        r"Technician|Representative|Designer|Architect|Host|Server|Bartender|"
        r"Cook|Chef|Cashier|Clerk|Teacher|Professor|Nurse|Doctor|Lawyer|"
        r"Accountant|Writer|Editor|Photographer|Artist|Scientist|Researcher))\s*$",
        re.IGNORECASE
    )
    
    # Normalize the text - collapse multiple newlines but preserve structure
    lines = experience_text.split("\n")
    
    # Group lines into job entries based on date patterns
    current_entry_lines = []
    entries = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if this line contains a date range (indicates new job entry)
        date_match = date_pattern.search(line)
        
        if date_match and current_entry_lines:
            # Save current entry and start a new one
            entries.append(current_entry_lines)
            current_entry_lines = [line]
        else:
            current_entry_lines.append(line)
    
    # Don't forget the last entry
    if current_entry_lines:
        entries.append(current_entry_lines)
    
    # If no date-based entries were found, try to group by double-spacing
    if len(entries) <= 1 and experience_text.strip():
        # Fallback: try splitting by multiple newlines
        paragraphs = re.split(r"\n\s*\n", experience_text)
        entries = [[p.strip()] for p in paragraphs if p.strip() and len(p.strip()) > 20]
    
    # Parse each entry
    for entry_lines in entries:
        if not entry_lines:
            continue
        
        raw_text = "\n".join(entry_lines)
        if len(raw_text) < 10:  # Skip very short entries
            continue
        
        entry = {
            "raw_text": raw_text,
            "company_name": None,
            "job_title": None,
            "start_date": None,
            "end_date": None,
            "is_current": False,
        }
        
        # Track which lines have been used for which purpose
        company_line_idx = -1
        
        # Try to extract structured fields from entry lines
        for idx, line in enumerate(entry_lines):
            line_stripped = line.strip()
            
            # Extract dates first (they're on the same line as job title usually)
            date_match = date_pattern.search(line_stripped)
            if date_match and not entry["start_date"]:
                date_str = date_match.group(0)
                entry["is_current"] = "present" in date_str.lower() or "current" in date_str.lower()
                # Parse date range
                dates = re.split(r"\s*[-–—to]+\s*", date_str, maxsplit=1)
                if len(dates) >= 1:
                    entry["start_date"] = _parse_date_string(dates[0])
                if len(dates) >= 2 and not entry["is_current"]:
                    entry["end_date"] = _parse_date_string(dates[1])
                
                # The remaining text after removing date is likely the job title
                line_no_date = date_pattern.sub("", line_stripped).strip()
                if line_no_date and not entry["job_title"]:
                    entry["job_title"] = line_no_date
            
            # Extract company and location (pattern: "Company Name – City, STATE")
            company_match = company_location_pattern.match(line_stripped)
            if company_match and not entry["company_name"]:
                entry["company_name"] = company_match.group(1).strip()
                company_line_idx = idx
        
        # If no job title found yet, look for title pattern in lines that aren't company
        if not entry["job_title"]:
            for idx, line in enumerate(entry_lines):
                if idx == company_line_idx:
                    continue  # Skip the company line
                line_stripped = line.strip()
                title_match = job_title_pattern.match(line_stripped)
                if title_match:
                    entry["job_title"] = title_match.group(1).strip()
                    break
        
        experiences.append(entry)
    
    return experiences[:20]  # Limit to 20 experiences


def _parse_date_string(date_str: str) -> date | None:
    """
    Parse a date string into a date object.
    
    Handles formats like: 01/2015, Jan 2015, January 2015, 2015
    """
    if not date_str:
        return None
    
    date_str = date_str.strip()
    
    # Skip "Present" or "Current"
    if date_str.lower() in ("present", "current"):
        return None
    
    try:
        # Try MM/YYYY format
        match = re.match(r"(\d{1,2})/(\d{4})", date_str)
        if match:
            month, year = int(match.group(1)), int(match.group(2))
            return date(year, month, 1)
        
        # Try Month YYYY format
        month_names = {
            "jan": 1, "january": 1, "feb": 2, "february": 2, "mar": 3, "march": 3,
            "apr": 4, "april": 4, "may": 5, "jun": 6, "june": 6,
            "jul": 7, "july": 7, "aug": 8, "august": 8, "sep": 9, "september": 9,
            "oct": 10, "october": 10, "nov": 11, "november": 11, "dec": 12, "december": 12
        }
        match = re.match(r"([A-Za-z]+)\s*(\d{4})", date_str)
        if match:
            month_str = match.group(1).lower()
            year = int(match.group(2))
            if month_str in month_names:
                return date(year, month_names[month_str], 1)
        
        # Try just YYYY
        match = re.match(r"^(\d{4})$", date_str)
        if match:
            return date(int(match.group(1)), 1, 1)
    except (ValueError, TypeError):
        pass
    
    return None


def _parse_education_section(education_text: str) -> list[dict[str, Any]]:
    """
    Parse education section into structured entries.
    
    Returns a list of education dictionaries with:
    - institution_name, degree_type, field_of_study, start_date, end_date, is_current
    """
    if not education_text:
        return []
    
    education = []
    
    # Common degree patterns - expanded to capture more variations
    # Use word boundaries and avoid matching state abbreviations like "MA" at end of addresses
    degree_pattern = re.compile(
        r"(?<![,\s][A-Z])"  # Negative lookbehind to avoid ", MA" (state abbrev)
        r"(?:"
        r"Bachelor(?:'s)?(?:\s+(?:of\s+)?(?:Science|Arts|Engineering|Business|Fine Arts))?"
        r"|Master(?:'s)?(?:\s+(?:of\s+)?(?:Science|Arts|Business|Engineering|Fine Arts))?"
        r"|PhD|Ph\.?D\.?"
        r"|Doctorate"
        r"|Associate(?:'s)?(?:\s+(?:of\s+)?(?:Science|Arts|Applied Science))?"
        r"|B\.S\.?(?:\s|$)|B\.A\.?(?:\s|$)"  # Require space or end after B.S./B.A.
        r"|M\.S\.?(?:\s|$)|M\.A\.?(?:\s|$)"  # Require space or end after M.S./M.A.
        r"|MBA|M\.B\.A\.?"
        r"|B\.?Sc\.?|M\.?Sc\.?"
        r"|BBA|B\.B\.A\.?"
        r"|Doctor of"
        r"|Diploma"
        r"|Certificate"
        r"|GED"
        r"|High School(?:\s+Diploma)?"
        r")"
        r"(?:\s+Degree)?",
        re.IGNORECASE
    )
    
    # Date pattern for graduation dates
    date_pattern = re.compile(
        r"(?:"
        # Date range: 2016-2020, 2016 - 2020
        r"(\d{4})\s*[-–—to]+\s*(\d{4}|[Pp]resent|[Cc]urrent|[Ee]xpected)"
        r"|"
        # Month Year format: May 2020, Expected May 2021
        r"(?:[Ee]xpected\s+)?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
        r"\s+(\d{4})"
        r"|"
        # Class of 2020
        r"Class of\s+(\d{4})"
        r")",
        re.IGNORECASE
    )
    
    # Institution pattern - universities, colleges, schools
    institution_pattern = re.compile(
        r"(?:University|College|Institute|School|Academy|Polytechnic|"
        r"Community College|State University|Technical School)",
        re.IGNORECASE
    )
    
    lines = education_text.split("\n")
    
    # Group lines into education entries based on degree patterns
    # Key insight: a new degree line typically starts a new entry
    current_entry_lines = []
    entries = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if this line starts a new education entry
        has_degree = degree_pattern.search(line)
        
        # A new degree line signals a new entry (if we already have content)
        if has_degree and current_entry_lines:
            # Save current entry and start a new one
            entries.append(current_entry_lines)
            current_entry_lines = [line]
        else:
            current_entry_lines.append(line)
    
    # Don't forget the last entry
    if current_entry_lines:
        entries.append(current_entry_lines)
    
    # If no entries were properly grouped (no degree patterns found), try fallback strategies
    if len(entries) <= 1 and education_text.strip():
        # Try splitting by double newlines
        paragraphs = re.split(r"\n\s*\n", education_text)
        if len(paragraphs) > 1:
            entries = [[p.strip()] for p in paragraphs if p.strip() and len(p.strip()) > 10]
        # Or try to identify entries by degree patterns
        elif len(entries) == 1:
            # Keep as is but try to parse better
            pass
    
    # Parse each entry
    for entry_lines in entries:
        if not entry_lines:
            continue
        
        raw_text = "\n".join(entry_lines)
        if len(raw_text) < 5:  # Skip very short entries
            continue
        
        entry = {
            "raw_text": raw_text,
            "institution_name": None,
            "degree_type": None,
            "field_of_study": None,
            "start_date": None,
            "end_date": None,
            "is_current": False,
        }
        
        # Process each line separately for better extraction
        institution_line_idx = -1
        degree_line_idx = -1
        
        for idx, line in enumerate(entry_lines):
            line = line.strip()
            
            # Check for institution on this line
            if institution_pattern.search(line) and not entry["institution_name"]:
                # Clean up institution name
                inst_name = line
                inst_name = date_pattern.sub("", inst_name)  # Remove dates
                inst_name = re.sub(r"\s*[-–—]+\s*[A-Za-z\s]+,\s*[A-Z]{2}\s*$", "", inst_name)  # Remove location suffix
                inst_name = inst_name.strip()
                if inst_name and len(inst_name) > 3:
                    entry["institution_name"] = inst_name
                    institution_line_idx = idx
            
            # Check for degree on this line
            degree_match = degree_pattern.search(line)
            if degree_match and not entry["degree_type"]:
                entry["degree_type"] = degree_match.group(0).strip()
                degree_line_idx = idx
                
                # Try to extract field of study from degree line
                # Pattern: "Bachelor of Science in Computer Science" -> "Computer Science"
                # Or: "BS in Computer Science" -> "Computer Science"
                after_degree = line[degree_match.end():].strip()
                # Also check before the degree for "in" patterns
                full_line_field_match = re.search(
                    r"(?:Bachelor|Master|Associate|Doctor)(?:'s)?\s+(?:of\s+)?(?:Science|Arts|Engineering|Business)?\s*(?:in|of)?\s+([A-Za-z\s,&]+?)(?:\s*[-–—,]|\s*$|\s*\d)",
                    line, re.IGNORECASE
                )
                if full_line_field_match:
                    field = full_line_field_match.group(1).strip()
                    field = re.sub(r"[,\.\-]+$", "", field).strip()
                    if field and len(field) > 2 and len(field) < 80:
                        entry["field_of_study"] = field
                elif "in " in after_degree.lower():
                    # Extract text after "in"
                    in_match = re.search(r"in\s+([A-Za-z\s,&]+?)(?:\s*[-–—,]|\s*$|\s*\d)", after_degree, re.IGNORECASE)
                    if in_match:
                        field = in_match.group(1).strip()
                        field = re.sub(r"[,\.\-]+$", "", field).strip()
                        if field and len(field) > 2 and len(field) < 80:
                            entry["field_of_study"] = field
        
        # If no institution found by pattern, look for lines with location patterns
        if not entry["institution_name"]:
            for idx, line in enumerate(entry_lines):
                line = line.strip()
                # Check for "Name – City, STATE" pattern
                loc_match = re.match(r"^(.+?)\s*[-–—]+\s*[A-Za-z\s]+,\s*[A-Z]{2}\s*$", line)
                if loc_match:
                    entry["institution_name"] = loc_match.group(1).strip()
                    break
        
        # Extract dates
        combined_text = " ".join(entry_lines)
        entry["is_current"] = bool(re.search(r"present|current|expected", combined_text, re.IGNORECASE))
        
        # Try to extract year(s)
        all_years = re.findall(r"\b(20\d{2}|19\d{2})\b", combined_text)
        if all_years:
            years = sorted([int(y) for y in set(all_years)])
            if len(years) >= 2:
                entry["start_date"] = date(years[0], 1, 1)
                if not entry["is_current"]:
                    entry["end_date"] = date(years[-1], 1, 1)
            elif len(years) == 1:
                # Single year - likely graduation year
                entry["end_date"] = date(years[0], 1, 1)
        
        education.append(entry)
    
    return education[:10]  # Limit to 10 education entries


def extract_structured_data(raw_text: str, file_type: str = "") -> dict[str, Any]:
    """
    Extract structured information from raw resume text using regex-based parsing.
    
    This implements Task 4268: Basic sectioning for Education, Experience, Skills.
    
    Args:
        raw_text: The raw text extracted from the resume file
        file_type: The file extension (e.g., ".pdf", ".docx") - for logging
        
    Returns:
        dict: Structured resume data containing:
            - full_name: Candidate's name
            - email: Email address
            - phone: Phone number
            - linkedin_url: LinkedIn profile URL
            - github_url: GitHub profile URL  
            - portfolio_url: Portfolio/website URL
            - professional_summary: Summary/objective section
            - skills: List of skills
            - experiences: List of work experience entries
            - education: List of education entries
            - raw_text: The original raw text for reference
    """
    log_info(f"Extracting structured data from {file_type} resume text...")
    
    # Initialize result structure
    result: dict[str, Any] = {
        "full_name": None,
        "email": None,
        "phone": None,
        "location": None,
        "linkedin_url": None,
        "github_url": None,
        "portfolio_url": None,
        "professional_summary": None,
        "skills": [],
        "experiences": [],
        "education": [],
        "certifications": [],
        "projects": [],
        "raw_text": raw_text,
    }
    
    if not raw_text or not raw_text.strip():
        log_warning("Empty text provided for extraction")
        return result
    
    # Extract contact information
    contact_info = _extract_contact_info(raw_text)
    result.update(contact_info)
    
    # Extract name
    result["full_name"] = _extract_name(raw_text)
    
    # Extract summary/objective section
    summary_content = _extract_section_content(raw_text, "summary")
    if summary_content:
        # Limit summary length
        result["professional_summary"] = summary_content[:2000]
    
    # Extract skills
    skills_content = _extract_section_content(raw_text, "skills")
    result["skills"] = _parse_skills_section(skills_content)
    
    # Extract experience
    experience_content = _extract_section_content(raw_text, "experience")
    result["experiences"] = _parse_experience_section(experience_content)
    
    # Extract education
    education_content = _extract_section_content(raw_text, "education")
    result["education"] = _parse_education_section(education_content)
    
    # Extract certifications (raw for now)
    cert_content = _extract_section_content(raw_text, "certifications")
    if cert_content:
        result["certifications"] = [{"raw_text": cert_content}]
    
    # Extract projects (raw for now)
    projects_content = _extract_section_content(raw_text, "projects")
    if projects_content:
        result["projects"] = [{"raw_text": projects_content}]
    
    log_info(
        f"Extraction complete: name={result['full_name']}, "
        f"email={result['email']}, skills={len(result['skills'])}, "
        f"experiences={len(result['experiences'])}, education={len(result['education'])}"
    )
    
    return result


# =============================================================================
# FILE TYPE DETECTION & ROUTING
# =============================================================================

def _get_file_extension(file_path: str) -> str:
    """Extract and validate file extension from path."""
    if "." not in file_path:
        return ""
    return "." + file_path.rsplit(".", 1)[-1].lower()


def parse_resume_content(file_content: bytes, file_path: str) -> dict[str, Any]:
    """
    Main entry point for parsing resume content.
    
    Detects file type and routes to appropriate parser.
    
    Args:
        file_content: The file content as bytes
        file_path: Original file path (used to determine file type)
        
    Returns:
        dict: Structured resume data
        
    Raises:
        ValueError: With specific error codes for parsing failures
    """
    file_ext = _get_file_extension(file_path)
    
    log_info(f"Parsing resume file: {file_path} (type: {file_ext})")
    
    # Validate file type
    if file_ext not in SUPPORTED_EXTENSIONS:
        log_error(f"Unsupported file type: {file_ext}")
        raise ValueError("ParseError: UnsupportedFileType")
    
    # Route to appropriate parser
    if file_ext == ".pdf":
        raw_text = parse_pdf_resume(file_content)
    elif file_ext == ".docx":
        raw_text = parse_docx_resume(file_content)
    else:
        raise ValueError("ParseError: UnsupportedFileType")
    
    # Extract structured data
    structured_data = extract_structured_data(raw_text, file_ext)
    
    return structured_data


# =============================================================================
# BACKGROUND TASK
# =============================================================================


def parse_resume_background(resume_id: str, file_path: str) -> None:
    """
    Background task to parse resume file and extract information.
    
    This function runs asynchronously after file upload completes.
    It downloads the file, parses content, extracts structured data,
    and updates the resume record with the extracted information.
    
    Args:
        resume_id: String UUID of the resume to parse
        file_path: Path to the uploaded resume file in storage
    
    Error Handling:
        - ParseError: ScannedPdfNoText - PDF is image-based/scanned
        - ParseError: EncryptedPdf - PDF is password protected
        - ParseError: CorruptedPdf - PDF file is corrupted
        - ParseError: CorruptedDocx - DOCX file is corrupted
        - ParseError: UnsupportedFileType - Unknown file extension
        - ParseError: Oversize - File exceeds 5MB limit
        - ParseError: FileNotFound - File not in storage
        - ParseError: DownloadFailed - Storage download error
    """
    db = SessionLocal()
    resume_uuid = UUID(resume_id)
    
    try:
        log_info(f"[RESUME_PARSE] Starting: resume_id={resume_id}")
        update_resume_status(resume_uuid, "Processing", db=db)
        
        statement = select(Resume).where(Resume.id == resume_uuid)
        resume = db.exec(statement).first()
        
        if not resume:
            raise ValueError(f"ParseError: ResumeNotFound - {resume_id}")
        
        log_info(f"[RESUME_PARSE] Downloading: resume_id={resume_id}, path={file_path}")
        file_content = download_resume_file(file_path)
        log_info(f"[RESUME_PARSE] Downloaded: size={len(file_content)} bytes")
        
        parsed_data = parse_resume_content(file_content, file_path)
        update_resume_with_parsed_data(resume, parsed_data, db)
        
        log_info(
            f"[RESUME_PARSE] Completed: resume_id={resume_id}, "
            f"name={parsed_data.get('full_name')}, "
            f"email={parsed_data.get('email')}, "
            f"skills={len(parsed_data.get('skills', []))}, "
            f"experience={len(parsed_data.get('experiences', []))}"
        )
        
        update_resume_status(resume_uuid, "Completed", db=db)
        
    except ValueError as e:
        error_msg = str(e)
        log_error(f"[RESUME_PARSE] Error: resume_id={resume_id}, error={error_msg}")
        update_resume_status(resume_uuid, "Failed", error_message=error_msg, db=db)
        
    except Exception as e:
        error_msg = f"ParseError: UnexpectedError - {str(e)}"
        log_error(f"[RESUME_PARSE] Unexpected: resume_id={resume_id}, error={error_msg}")
        update_resume_status(resume_uuid, "Failed", error_message=error_msg, db=db)
        
    finally:
        db.close()


def get_active_resume(
    user_id: UUID,
    db: Session
) -> Resume | None:
    """
    Retrieve the most recently uploaded active resume for the user.
    
    Args:
        user_id: UUID of the current user
        db: Database session
    Returns:
        Resume instance if found, None otherwise
    """
    statement = select(Resume).where(
        Resume.user_id == user_id,
        Resume.deleted_at.is_(None)
    ).order_by(Resume.created_at.desc())
    
    result = db.exec(statement)
    resume = result.first()
    
    return resume


def get_resume_complete(
    resume_id: UUID,
    user_id: UUID,
    db: Session
) -> dict | None:
    """
    Retrieve a resume with all related sections (experiences, education, skills, etc.).
    
    This function fetches the base resume and all associated structured data
    from the normalized tables, returning a complete picture for the frontend.
    
    Args:
        resume_id: UUID of the resume to retrieve
        user_id: UUID of the user (for authorization check)
        db: Database session
        
    Returns:
        dict containing resume data with all sections, or None if not found/unauthorized
    """
    # Fetch base resume with authorization check
    statement = select(Resume).where(
        Resume.id == resume_id,
        Resume.user_id == user_id
    )
    resume = db.exec(statement).first()
    
    if not resume:
        return None
    
    # Fetch all related sections
    experiences = db.exec(
        select(ResumeExperience)
        .where(ResumeExperience.resume_id == resume_id)
        .order_by(ResumeExperience.display_order)
    ).all()
    
    education = db.exec(
        select(ResumeEducation)
        .where(ResumeEducation.resume_id == resume_id)
        .order_by(ResumeEducation.display_order)
    ).all()
    
    skills = db.exec(
        select(ResumeSkill)
        .where(ResumeSkill.resume_id == resume_id)
        .order_by(ResumeSkill.display_order)
    ).all()
    
    certifications = db.exec(
        select(ResumeCertification)
        .where(ResumeCertification.resume_id == resume_id)
        .order_by(ResumeCertification.display_order)
    ).all()
    
    projects = db.exec(
        select(ResumeProject)
        .where(ResumeProject.resume_id == resume_id)
        .order_by(ResumeProject.display_order)
    ).all()
    
    log_info(
        f"Fetched complete resume {resume_id}: "
        f"experiences={len(experiences)}, education={len(education)}, "
        f"skills={len(skills)}, certifications={len(certifications)}, "
        f"projects={len(projects)}"
    )
    
    return {
        "resume": resume,
        "experiences": experiences,
        "education": education,
        "skills": skills,
        "certifications": certifications,
        "projects": projects,
    }


def duplicate_resume(
    resume_id: UUID,
    user_id: UUID,
    db: Session,
    new_version_name: str | None = None
) -> Resume | None:
    """
    Duplicate a resume with all its sections.
    
    Creates a complete copy of the resume including all experiences,
    education, skills, projects, and certifications with new UUIDs.
    
    Args:
        resume_id: ID of the resume to duplicate
        user_id: ID of the authenticated user
        db: Database session
        new_version_name: Optional custom name for the duplicate
        
    Returns:
        New Resume instance if duplication successful, None if original not found
        
    Raises:
        Exception: If database operations fail
    """
    # Get the complete resume data
    complete_data = get_resume_complete(resume_id, user_id, db)
    if not complete_data:
        return None
    
    original_resume = complete_data["resume"]
    
    # Generate unique version name
    if not new_version_name:
        base_name = original_resume.version_name
        # Find existing duplicates to avoid name conflicts
        existing_stmt = select(Resume).where(
            Resume.user_id == user_id,
            Resume.version_name.like(f"{base_name} (Copy%)")
        )
        existing_copies = db.exec(existing_stmt).all()
        copy_number = len(existing_copies) + 1
        new_version_name = f"{base_name} (Copy {copy_number})"
    
    try:
        # Create the new resume
        new_resume = Resume(
            user_id=user_id,
            version_name=new_version_name,
            template_id=original_resume.template_id,
            is_primary=False,  # Duplicates are never primary
            section_order=original_resume.section_order,
            file_path=None,  # Don't copy file references
            file_url=None,
            full_name=original_resume.full_name,
            email=original_resume.email,
            phone=original_resume.phone,
            location=original_resume.location,
            linkedin_url=original_resume.linkedin_url,
            github_url=original_resume.github_url,
            portfolio_url=original_resume.portfolio_url,
            professional_summary=original_resume.professional_summary,
            raw_text=None,  # Don't copy raw parsing data
            processing_status="Completed",  # Duplicates are immediately completed
            error_message=None
        )
        
        db.add(new_resume)
        db.commit()
        db.refresh(new_resume)
        
        # Copy all experiences
        for exp in complete_data["experiences"]:
            new_exp = ResumeExperience(
                resume_id=new_resume.id,
                company_name=exp.company_name,
                job_title=exp.job_title,
                location=exp.location,
                start_date=exp.start_date,
                end_date=exp.end_date,
                is_current=exp.is_current,
                description=exp.description,
                achievements=exp.achievements,
                skills_used=exp.skills_used,
                display_order=exp.display_order
            )
            db.add(new_exp)
        
        # Copy all education
        for edu in complete_data["education"]:
            new_edu = ResumeEducation(
                resume_id=new_resume.id,
                institution_name=edu.institution_name,
                degree_type=edu.degree_type,
                field_of_study=edu.field_of_study,
                location=edu.location,
                start_date=edu.start_date,
                end_date=edu.end_date,
                is_current=edu.is_current,
                gpa=edu.gpa,
                achievements=edu.achievements,
                relevant_coursework=edu.relevant_coursework,
                display_order=edu.display_order
            )
            db.add(new_edu)
        
        # Copy all skills
        for skill in complete_data["skills"]:
            new_skill = ResumeSkill(
                resume_id=new_resume.id,
                skill_name=skill.skill_name,
                skill_category=skill.skill_category,
                proficiency_level=skill.proficiency_level,
                years_of_experience=skill.years_of_experience,
                is_primary=skill.is_primary,
                display_order=skill.display_order
            )
            db.add(new_skill)
        
        # Copy all projects
        for project in complete_data["projects"]:
            new_project = ResumeProject(
                resume_id=new_resume.id,
                project_name=project.project_name,
                role=project.role,
                description=project.description,
                technologies_used=project.technologies_used,
                project_url=project.project_url,
                start_date=project.start_date,
                end_date=project.end_date,
                is_current=project.is_current,
                achievements=project.achievements,
                display_order=project.display_order
            )
            db.add(new_project)
        
        # Copy all certifications
        for cert in complete_data["certifications"]:
            new_cert = ResumeCertification(
                resume_id=new_resume.id,
                certification_name=cert.certification_name,
                issuing_organization=cert.issuing_organization,
                issue_date=cert.issue_date,
                expiry_date=cert.expiry_date,
                credential_id=cert.credential_id,
                credential_url=cert.credential_url,
                display_order=cert.display_order
            )
            db.add(new_cert)
        
        db.commit()
        
        log_info(f"Resume duplicated successfully: original_id={resume_id}, new_id={new_resume.id}, user_id={user_id}")
        return new_resume
        
    except Exception as e:
        db.rollback()
        log_error(f"Failed to duplicate resume {resume_id}: {str(e)}")
        raise